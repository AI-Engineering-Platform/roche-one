import json
from pathlib import Path
from docx import Document
from PyPDF2 import PdfReader


def read_json(path: str) -> dict:
    p = Path(path)
    with p.open("r", encoding="utf-8") as f:
        return json.load(f)


def read_docx_text(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def read_pdf_text(path: str) -> str:
    reader = PdfReader(path)
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text)
    return "\n".join(pages_text)


def write_docx_text(path: str, text: str) -> None:
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)
